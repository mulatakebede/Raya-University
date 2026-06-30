# -*- coding: utf-8 -*-
"""
Raya University - Liaison Office Guest House
Guest Satisfaction Survey Analysis 2018
"""

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# ============================================================
# RAYA UNIVERSITY BRANDING
# ============================================================

UNIVERSITY_NAME = "Raya University"
DEPARTMENT = "Liaison Office Guest House"
SURVEY_YEAR = "2018"
LOGO_TEXT = "RU"  # Placeholder for logo

# Raya University Colors
RAYA_NAVY = "#1A2A5E"      # Primary Navy Blue
RAYA_GOLD = "#C89B3C"       # Gold/Amber
RAYA_LIGHT_BLUE = "#4A7FC1" # Secondary Blue
RAYA_CREAM = "#F5F0E6"      # Cream/Off-white
RAYA_DARK = "#0D1A3A"       # Dark Navy

# Color palette for charts
NAVY = RAYA_NAVY
BLUE = RAYA_LIGHT_BLUE
LIGHT_BLUE = "#9DC3E6"
GREEN = "#63BE7B"
LIGHT_GREEN = "#A9D18E"
YELLOW = "#FFEB84"
ORANGE = "#F4B183"
RED = "#E06666"
GREY = "#595959"
DARK_GREY = "#333333"
GOLD = RAYA_GOLD

plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['font.size'] = 10
plt.rcParams['figure.dpi'] = 150
plt.rcParams['axes.edgecolor'] = RAYA_NAVY
plt.rcParams['axes.linewidth'] = 0.8

# ============================================================
# MAPPINGS
# ============================================================

SATISFACTION_MAP = {'ሀ': 5, 'ለ': 4, 'ሐ': 3, 'መ': 2, 'ሠ': 1, 'ሰ': 1}
SATISFACTION_LABEL_MAP = {
    'ሀ': 'Very Satisfied', 'ለ': 'Satisfied', 'ሐ': 'Neutral',
    'መ': 'Dissatisfied', 'ሠ': 'Very Dissatisfied', 'ሰ': 'Very Dissatisfied'
}

QUALITY_MAP = {'ሀ': 1, 'ለ': 2, 'ሐ': 3, 'መ': 4, 'ሠ': 5, 'ሰ': 5}
QUALITY_LABEL_MAP = {
    'ሀ': 'Very Low', 'ለ': 'Low', 'ሐ': 'Medium', 'መ': 'Good',
    'ሠ': 'Very Good', 'ሰ': 'Very Good'
}

PURPOSE_MAP = {'ሀ': 'Work', 'ለ': 'Training', 'ሐ': 'Meeting', 'መ': 'Other'}
DURATION_MAP = {'ሀ': '1 day', 'ለ': '2-3 days', 'ሐ': 'More than 3 days'}
YES_NO_MAP = {'ሀ': 'Yes', 'ለ': 'No', 'ሐ': 'Maybe'}

COLUMN_NAMES = [
    'respondent_id', 'visit_purpose', 'stay_duration',
    'q1_reception_speed', 'q2_staff_respect', 'q3_info_clarity', 'q4_hospitality',
    'q5_room_cleanliness', 'q6_bed_quality', 'q7_water_bathroom', 'q8_electricity_basic',
    'q9_environment_safety', 'q10_cleaning_service', 'q11_wifi_communication',
    'q12_problem_resolution', 'q13_overall_service', 'q14_return_willingness',
    'q15_recommend', 'q16_liaison_staff_capacity', 'open_1', 'open_2', 'open_3'
]

QUESTION_LABELS = {
    'q1_reception_speed': 'Reception Speed',
    'q2_staff_respect': 'Staff Respect & Courtesy',
    'q3_info_clarity': 'Information Clarity',
    'q4_hospitality': 'Hospitality',
    'q5_room_cleanliness': 'Room Cleanliness',
    'q6_bed_quality': 'Bed & Furniture Quality',
    'q7_water_bathroom': 'Water & Bathroom Service',
    'q8_electricity_basic': 'Electricity & Basic Services',
    'q9_environment_safety': 'Environmental Safety & Security',
    'q10_cleaning_service': 'Cleaning Service',
    'q11_wifi_communication': 'Wi-Fi & Communication',
    'q12_problem_resolution': 'Problem Resolution Speed',
    'q13_overall_service': 'Overall Service',
    'q16_liaison_staff_capacity': 'Liaison Staff Capacity'
}

# ============================================================
# DATA LOADING & PROCESSING
# ============================================================

def load_data():
    """Load data from Excel file"""
    data_file = None
    for f in os.listdir('.'):
        if f.endswith('.xlsx') and ('መልስ' in f or '2018' in f):
            data_file = f
            break
    
    if not data_file:
        print("ERROR: Could not find the Excel data file!")
        print("Looking for file containing 'መልስ' or '2018'")
        return None
    
    print(f"Loading data from: {data_file}")
    df = pd.read_excel(data_file, sheet_name='Sheet1', header=None, skiprows=2)
    n_cols = min(len(COLUMN_NAMES), len(df.columns))
    df.columns = COLUMN_NAMES[:n_cols]
    df = df.dropna(how='all').reset_index(drop=True)
    
    for col in df.select_dtypes(include=['object']).columns:
        if col in df.columns:
            df[col] = df[col].astype(str).str.strip().replace('nan', np.nan).replace('', np.nan)
    
    df = df[pd.to_numeric(df['respondent_id'], errors='coerce').notna()].reset_index(drop=True)
    print(f"Loaded {len(df)} respondents")
    return df

def process_data(df):
    """Process the survey data with CORRECT mappings"""
    print("Processing data...")
    
    def map_response(value, mapping):
        if pd.isna(value) or value == '' or value == 'nan':
            return np.nan
        val_str = str(value)
        if val_str and len(val_str) > 0:
            return mapping.get(val_str[0], np.nan)
        return np.nan
    
    sat_questions = [
        'q1_reception_speed', 'q2_staff_respect', 'q3_info_clarity', 'q4_hospitality',
        'q5_room_cleanliness', 'q6_bed_quality', 'q7_water_bathroom', 'q8_electricity_basic',
        'q9_environment_safety', 'q10_cleaning_service', 'q11_wifi_communication',
        'q12_problem_resolution'
    ]
    
    for col in sat_questions:
        if col in df.columns:
            df[f'{col}_num'] = df[col].apply(lambda x: map_response(x, SATISFACTION_MAP))
            df[f'{col}_label'] = df[col].apply(lambda x: map_response(x, SATISFACTION_LABEL_MAP))
    
    if 'q13_overall_service' in df.columns:
        df['q13_overall_service_num'] = df['q13_overall_service'].apply(lambda x: map_response(x, QUALITY_MAP))
        df['q13_overall_service_label'] = df['q13_overall_service'].apply(lambda x: map_response(x, QUALITY_LABEL_MAP))
    
    if 'q16_liaison_staff_capacity' in df.columns:
        df['q16_liaison_staff_capacity_num'] = df['q16_liaison_staff_capacity'].apply(lambda x: map_response(x, QUALITY_MAP))
        df['q16_liaison_staff_capacity_label'] = df['q16_liaison_staff_capacity'].apply(lambda x: map_response(x, QUALITY_LABEL_MAP))
    
    if 'visit_purpose' in df.columns:
        df['visit_purpose_label'] = df['visit_purpose'].apply(lambda x: map_response(x, PURPOSE_MAP))
    if 'stay_duration' in df.columns:
        df['stay_duration_label'] = df['stay_duration'].apply(lambda x: map_response(x, DURATION_MAP))
    if 'q14_return_willingness' in df.columns:
        df['q14_return_willingness_label'] = df['q14_return_willingness'].apply(lambda x: map_response(x, YES_NO_MAP))
    if 'q15_recommend' in df.columns:
        df['q15_recommend_label'] = df['q15_recommend'].apply(lambda x: 
            {'ሀ': 'Yes', 'ለ': 'No'}.get(str(x)[0] if pd.notna(x) else '', np.nan))
    
    print(f"Processed {len(df)} records")
    return df

# ============================================================
# CHART GENERATION WITH RAYA UNIVERSITY BRANDING
# ============================================================

def create_charts(df):
    """Generate all visualization charts with Raya University branding"""
    os.makedirs('charts', exist_ok=True)
    color_seq = [GREEN, LIGHT_GREEN, YELLOW, ORANGE, RED]
    
    # Chart 1: Mean Scores
    print("Creating chart_means.png...")
    means = []
    for col in [c for c in df.columns if c.endswith('_num')]:
        m = df[col].mean()
        if not pd.isna(m):
            label = QUESTION_LABELS.get(col.replace('_num', ''), col)
            means.append({'variable': label, 'mean': m})
    
    if means:
        df_m = pd.DataFrame(means).sort_values('mean', ascending=True)
        fig, ax = plt.subplots(figsize=(10, 7))
        colors = [GREEN if m >= 4.5 else LIGHT_GREEN if m >= 4.0 else ORANGE if m >= 3.5 else RED for m in df_m['mean']]
        bars = ax.barh(df_m['variable'], df_m['mean'], color=colors, height=0.6, edgecolor='white', linewidth=0.5)
        ax.set_xlim(1, 5)
        ax.set_xlabel('Mean Score (1 = Worst → 5 = Best)', fontsize=11, color=GREY)
        ax.set_title(f'{UNIVERSITY_NAME}\nMean Satisfaction Scores by Service Area', 
                    fontsize=14, fontweight='bold', color=RAYA_NAVY, pad=15)
        ax.axvline(3, color='#999999', linestyle='--', linewidth=1)
        ax.text(3, -0.5, 'neutral', fontsize=9, color='#888888', ha='center')
        for bar, m in zip(bars, df_m['mean']):
            ax.text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2, f'{m:.2f}', va='center', fontsize=9)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.tick_params(axis='y', labelsize=9)
        
        # Add Raya University watermark
        ax.text(0.98, 0.02, f'{UNIVERSITY_NAME}', transform=ax.transAxes, 
                fontsize=8, color='#cccccc', ha='right', va='bottom', alpha=0.5)
        
        plt.tight_layout()
        plt.savefig('charts/chart_means.png', dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
    
    # Chart 2: Visit Profile
    print("Creating chart_visit_profile.png...")
    fig, axes = plt.subplots(1, 2, figsize=(10, 4))
    fig.suptitle(f'{UNIVERSITY_NAME} - Guest Profile', fontsize=12, fontweight='bold', color=RAYA_NAVY)
    
    if 'visit_purpose_label' in df.columns:
        freq = df['visit_purpose_label'].value_counts()
        if len(freq) > 0:
            axes[0].bar(freq.index, freq.values, color=RAYA_NAVY, width=0.5)
            axes[0].set_title('Purpose of Visit', fontsize=12, fontweight='bold', color=RAYA_NAVY)
            for i, (label, count) in enumerate(freq.items()):
                if not pd.isna(label):
                    axes[0].text(i, count + 0.3, str(count), ha='center', fontsize=10, fontweight='bold')
            axes[0].spines['top'].set_visible(False)
            axes[0].spines['right'].set_visible(False)
            axes[0].set_ylabel('Number of Respondents', fontsize=10)
            axes[0].tick_params(axis='x', rotation=15, labelsize=9)
    
    if 'stay_duration_label' in df.columns:
        freq = df['stay_duration_label'].value_counts()
        if len(freq) > 0:
            axes[1].bar(freq.index, freq.values, color=RAYA_GOLD, width=0.5)
            axes[1].set_title('Length of Stay', fontsize=12, fontweight='bold', color=RAYA_NAVY)
            for i, (label, count) in enumerate(freq.items()):
                if not pd.isna(label):
                    axes[1].text(i, count + 0.3, str(count), ha='center', fontsize=10, fontweight='bold')
            axes[1].spines['top'].set_visible(False)
            axes[1].spines['right'].set_visible(False)
            axes[1].tick_params(axis='x', labelsize=9)
    
    plt.tight_layout()
    plt.savefig('charts/chart_visit_profile.png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    
    # Chart 3: Loyalty Indicators
    print("Creating chart_loyalty.png...")
    fig, ax = plt.subplots(figsize=(8, 4.5))
    
    r_dict = {'Yes': 0, 'No': 0, 'Maybe': 0}
    if 'q14_return_willingness_label' in df.columns:
        for label, count in df['q14_return_willingness_label'].value_counts().items():
            if not pd.isna(label) and label in r_dict:
                r_dict[label] = count
    
    rec_dict = {'Yes': 0, 'No': 0, 'Maybe': 0}
    if 'q15_recommend_label' in df.columns:
        for label, count in df['q15_recommend_label'].value_counts().items():
            if not pd.isna(label) and label in rec_dict:
                rec_dict[label] = count
    
    total_r = sum(r_dict.values()) or 1
    total_rec = sum(rec_dict.values()) or 1
    cats = ['Yes', 'No', 'Maybe']
    ret_vals = [r_dict[c]/total_r*100 for c in cats]
    rec_vals = [rec_dict[c]/total_rec*100 for c in cats]
    
    x = range(len(cats))
    width = 0.35
    bars1 = ax.bar([i - width/2 for i in x], ret_vals, width=width, label='Willing to Return', color=RAYA_NAVY)
    bars2 = ax.bar([i + width/2 for i in x], rec_vals, width=width, label='Would Recommend', color=RAYA_GOLD)
    
    ax.set_xticks(list(x))
    ax.set_xticklabels(cats, fontsize=10)
    ax.set_ylabel('% of Valid Responses', fontsize=10)
    ax.set_title(f'{UNIVERSITY_NAME}\nLoyalty Indicators', fontsize=13, fontweight='bold', color=RAYA_NAVY)
    
    for bars in (bars1, bars2):
        for bar in bars:
            h = bar.get_height()
            if h > 0:
                ax.text(bar.get_x() + bar.get_width()/2, h + 1.5, f'{h:.0f}%', ha='center', fontsize=9, fontweight='bold')
    
    ax.legend(frameon=False, fontsize=10)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_ylim(0, 110)
    
    plt.tight_layout()
    plt.savefig('charts/chart_loyalty.png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    
    # Chart 4: Distribution
    print("Creating chart_distribution.png...")
    focus_vars = ['q7_water_bathroom_num', 'q13_overall_service_num', 'q16_liaison_staff_capacity_num']
    focus_labels = ['Water & Bathroom Service', 'Overall Service', 'Liaison Staff Capacity']
    fig, axes = plt.subplots(len(focus_vars), 1, figsize=(9, 5.5))
    order = ['Very Good', 'Good', 'Medium', 'Low', 'Very Low']
    
    for idx, (ax, var, label) in enumerate(zip(axes, focus_vars, focus_labels)):
        if var in df.columns:
            freq = df[var].value_counts()
            vals = [freq.get(v, 0) for v in [5, 4, 3, 2, 1]]
            total = sum(vals) or 1
            pct_vals = [v/total*100 for v in vals]
            left = 0
            for val, pct, col in zip(order, pct_vals, color_seq):
                ax.barh([label], [pct], left=left, color=col, height=0.5, edgecolor='white', linewidth=0.5)
                if pct >= 3:
                    ax.text(left + pct/2, 0, f'{pct:.0f}%', ha='center', va='center', fontsize=8,
                           fontweight='bold', color='white' if pct > 30 else DARK_GREY)
                left += pct
            ax.set_xlim(0, 100)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_visible(False)
            ax.tick_params(axis='y', labelsize=10)
            ax.set_yticks([0])
            ax.set_yticklabels([label], fontsize=10, fontweight='bold')
    
    axes[-1].set_xlabel('% of Valid Responses', fontsize=10, color=GREY)
    fig.suptitle(f'{UNIVERSITY_NAME}\nDistribution of Ratings for Key Service Areas', 
                fontsize=13, fontweight='bold', color=RAYA_NAVY, y=0.98)
    
    legend_items = [Patch(facecolor=c, label=o) for c, o in zip(color_seq, order)]
    fig.legend(handles=legend_items, loc='lower center', ncol=5, frameon=False, fontsize=8, bbox_to_anchor=(0.5, -0.02))
    
    plt.tight_layout(rect=[0, 0.06, 1, 0.94])
    plt.savefig('charts/chart_distribution.png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    
    print("All charts created!")

# ============================================================
# EXCEL EXPORT
# ============================================================

def export_excel(df):
    """Export all analysis results to Excel"""
    print("Exporting survey_analysis_2018.xlsx...")
    with pd.ExcelWriter('survey_analysis_2018.xlsx', engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Raw Data', index=False)
        
        stats = []
        for col in [c for c in df.columns if c.endswith('_num')]:
            data = df[col].dropna()
            if len(data) > 0:
                label = QUESTION_LABELS.get(col.replace('_num', ''), col)
                stats.append({
                    'Service Area': label,
                    'N': len(data),
                    'Mean': round(data.mean(), 2),
                    'Median': round(data.median(), 2),
                    'Std Dev': round(data.std(), 2),
                    'Min': data.min(),
                    'Max': data.max()
                })
        if stats:
            pd.DataFrame(stats).to_excel(writer, sheet_name='Summary Stats', index=False)
        
        for col in [c for c in df.columns if c.endswith('_label')]:
            if col not in ['q14_return_willingness_label', 'q15_recommend_label']:
                freq = df[col].value_counts(dropna=False)
                valid = df[col].dropna().value_counts()
                total = len(df)
                valid_total = len(df[col].dropna()) or 1
                freq_df = pd.DataFrame({
                    'Response': freq.index,
                    'Count': freq.values,
                    'Percent': [round(v/total*100, 1) for v in freq.values],
                    'Valid %': [round(v/valid_total*100, 1) for v in 
                               [valid.get(x, 0) if not pd.isna(x) else 0 for x in freq.index]]
                })
                base = col.replace('_label', '')
                sheet = QUESTION_LABELS.get(base, base)[:31]
                freq_df.to_excel(writer, sheet_name=sheet, index=False)
        
        for col in ['visit_purpose_label', 'stay_duration_label']:
            if col in df.columns:
                freq = df[col].value_counts()
                pd.DataFrame({
                    'Category': freq.index,
                    'Count': freq.values,
                    'Percent': [round(v/len(df)*100, 1) for v in freq.values]
                }).to_excel(writer, sheet_name=col.replace('_label', '')[:31], index=False)
        
        loyalty = []
        for col in ['q14_return_willingness_label', 'q15_recommend_label']:
            if col in df.columns:
                for val, count in df[col].value_counts(dropna=False).items():
                    loyalty.append({
                        'Indicator': col.replace('_label', '').replace('_', ' ').title(),
                        'Response': val if not pd.isna(val) else 'Missing',
                        'Count': count,
                        'Percent': round(count/len(df)*100, 1)
                    })
        if loyalty:
            pd.DataFrame(loyalty).to_excel(writer, sheet_name='Loyalty', index=False)
    
    print("Excel export complete!")

# ============================================================
# WORD REPORT
# ============================================================

def generate_word_report(df):
    """Generate Word document report"""
    print("Generating survey_report_2018.docx...")
    doc = Document()
    
    # Title with Raya University branding
    title = doc.add_heading(f'{UNIVERSITY_NAME}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f'{DEPARTMENT}\nGuest Satisfaction Survey Report {SURVEY_YEAR}', 1)
    doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    n = len(df)
    all_means = [df[col].mean() for col in df.columns if col.endswith('_num') and not pd.isna(df[col].mean())]
    overall_mean = np.mean(all_means) if all_means else 0
    
    means_data = []
    for col in [c for c in df.columns if c.endswith('_num')]:
        m = df[col].mean()
        if not pd.isna(m):
            label = QUESTION_LABELS.get(col.replace('_num', ''), col)
            means_data.append({'label': label, 'mean': m})
    means_df = pd.DataFrame(means_data).sort_values('mean', ascending=False)
    top_3 = means_df.head(3)['label'].tolist()
    
    return_rate = round(df['q14_return_willingness_label'].value_counts().get('Yes', 0) / len(df['q14_return_willingness_label'].dropna()) * 100, 1) if 'q14_return_willingness_label' in df.columns else 0
    rec_rate = round(df['q15_recommend_label'].value_counts().get('Yes', 0) / len(df['q15_recommend_label'].dropna()) * 100, 1) if 'q15_recommend_label' in df.columns else 0
    
    exec_summary = f"""
A guest satisfaction survey was conducted at the {DEPARTMENT} to assess the quality of service provided to visitors. A total of {n} guests completed the survey, rating their experience across reception, accommodation, support services, and overall satisfaction.

Overall, guests reported a high level of satisfaction, with an average score of {overall_mean:.2f} out of 5 across all measured service dimensions. The strongest-performing areas were {top_3[0] if len(top_3)>0 else 'N/A'}, {top_3[1] if len(top_3)>1 else 'N/A'}, and {top_3[2] if len(top_3)>2 else 'N/A'}, while water and bathroom service received the most varied feedback and represents the clearest opportunity for improvement.

{return_rate}% of respondents indicated they would return to the guesthouse, and {rec_rate}% said they would recommend it to others — both strong indicators of guest loyalty and satisfaction.
"""
    doc.add_paragraph(exec_summary.strip())
    doc.add_page_break()
    
    # Section 2: Methodology
    doc.add_heading('2. Methodology', level=1)
    methodology = f"""
Data was collected through a structured, bilingual (Amharic/English) paper-based questionnaire administered to guests of the {DEPARTMENT}. Responses were later digitized and analyzed.

• Sample size: {n} respondents
• Survey instrument: 18-item structured questionnaire across six sections
• Rating scale: 5-point Likert scale (1 = Very dissatisfied/Very low, 5 = Very satisfied/Very good)
• Data collection method: Manual paper survey, later digitized for analysis
"""
    doc.add_paragraph(methodology.strip())
    doc.add_page_break()
    
    # Section 3: Respondent Profile
    doc.add_heading('3. Respondent Profile', level=1)
    if 'visit_purpose_label' in df.columns:
        doc.add_heading('3.1 Purpose of Visit', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Purpose'; hdr[1].text = 'Count'; hdr[2].text = 'Percent'
        for purpose, count in df['visit_purpose_label'].value_counts().items():
            if not pd.isna(purpose):
                row = table.add_row().cells
                row[0].text = str(purpose)
                row[1].text = str(count)
                row[2].text = f'{round(count/len(df)*100, 1)}%'
    
    if 'stay_duration_label' in df.columns:
        doc.add_heading('3.2 Length of Stay', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Duration'; hdr[1].text = 'Count'; hdr[2].text = 'Percent'
        for duration, count in df['stay_duration_label'].value_counts().items():
            if not pd.isna(duration):
                row = table.add_row().cells
                row[0].text = str(duration)
                row[1].text = str(count)
                row[2].text = f'{round(count/len(df)*100, 1)}%'
    doc.add_page_break()
    
    # Section 4: Overall Satisfaction
    doc.add_heading('4. Overall Satisfaction', level=1)
    if 'q13_overall_service_num' in df.columns:
        overall_score = df['q13_overall_service_num'].mean()
        doc.add_paragraph(f"The overall service satisfaction score averages {overall_score:.2f} out of 5.0.")
        doc.add_heading('4.1 Distribution of Overall Service Ratings', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Rating'; hdr[1].text = 'Count'; hdr[2].text = 'Percent'
        for rating in ['Very Good', 'Good', 'Medium', 'Low', 'Very Low']:
            count = df['q13_overall_service_label'].value_counts().get(rating, 0)
            if count > 0:
                row = table.add_row().cells
                row[0].text = rating
                row[1].text = str(count)
                row[2].text = f'{round(count/len(df["q13_overall_service_label"].dropna())*100, 1)}%'
    doc.add_page_break()
    
    # Section 5: Service Area Analysis
    doc.add_heading('5. Service Area Analysis', level=1)
    doc.add_heading('5.1 Mean Scores by Service Area', level=2)
    stats = []
    for col in [c for c in df.columns if c.endswith('_num')]:
        data = df[col].dropna()
        if len(data) > 0:
            label = QUESTION_LABELS.get(col.replace('_num', ''), col)
            stats.append({
                'Service Area': label,
                'Mean': round(data.mean(), 2),
                'Median': round(data.median(), 2),
                'Std Dev': round(data.std(), 2),
                'N': len(data)
            })
    
    if stats:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        headers = ['Service Area', 'Mean', 'Median', 'Std Dev', 'N']
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for s in sorted(stats, key=lambda x: x['Mean'], reverse=True):
            row = table.add_row().cells
            row[0].text = s['Service Area']
            row[1].text = f"{s['Mean']:.2f}"
            row[2].text = f"{s['Median']:.1f}"
            row[3].text = f"{s['Std Dev']:.2f}"
            row[4].text = str(s['N'])
    doc.add_page_break()
    
    # Section 6: Loyalty Indicators
    doc.add_heading('6. Loyalty Indicators', level=1)
    for col, title in [('q14_return_willingness_label', 'Willingness to Return'),
                       ('q15_recommend_label', 'Willingness to Recommend')]:
        if col in df.columns:
            idx = ['q14_return_willingness_label', 'q15_recommend_label'].index(col)
            doc.add_heading(f'6.{["a","b"][idx]} {title}', level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Response'; hdr[1].text = 'Count'; hdr[2].text = 'Percent'
            for response, count in df[col].value_counts().items():
                if not pd.isna(response):
                    row = table.add_row().cells
                    row[0].text = str(response)
                    row[1].text = str(count)
                    row[2].text = f'{round(count/len(df[col].dropna())*100, 1)}%'
    doc.add_page_break()
    
    # Section 7: Recommendations
    doc.add_heading('7. Recommendations', level=1)
    if stats:
        stats_df = pd.DataFrame(stats).sort_values('Mean', ascending=False)
        top_areas = stats_df.head(3)
        
        recommendations = f"""
Based on the survey findings, the following recommendations are proposed:

1. IMPROVE WATER AND BATHROOM FACILITIES (Priority #1)
   • This area received the lowest satisfaction score
   • Consider upgrading plumbing infrastructure
   • Ensure consistent hot water supply
   • Improve water pressure and bathroom cleanliness

2. ENHANCE Wi-Fi AND COMMUNICATION SERVICES (Priority #2)
   • Second lowest satisfaction area
   • Invest in better internet infrastructure
   • Ensure reliable connectivity throughout the facility

3. IMPROVE PROBLEM RESOLUTION SPEED
   • Guests reported delays in getting issues addressed
   • Implement a faster response system
   • Train staff on urgent issue handling

4. ADDRESS ELECTRICITY AND BASIC SERVICES
   • Ensure consistent power supply
   • Install backup generators for outages
   • Regular maintenance of electrical systems

5. MAINTAIN STRENGTHS (High-Performing Areas)
   • {top_areas.iloc[0]['Service Area'] if len(top_areas) > 0 else 'Environmental Safety'} scored highest
   • Continue current practices in high-performing areas
   • Maintain staff training programs
   • Implement regular guest feedback to monitor ongoing satisfaction
"""
        doc.add_paragraph(recommendations.strip())
    
    doc.save('survey_report_2018.docx')
    print("Word report complete!")

# ============================================================
# WEB DASHBOARD WITH RAYA UNIVERSITY BRANDING
# ============================================================

def generate_dashboard(df):
    """Generate complete HTML dashboard with Raya University branding"""
    print("Generating dashboard.html...")
    
    # Calculate metrics
    n = len(df)
    all_means = [df[col].mean() for col in df.columns if col.endswith('_num') and not pd.isna(df[col].mean())]
    overall_mean = np.mean(all_means) if all_means else 0
    
    means_data = []
    for col in [c for c in df.columns if c.endswith('_num')]:
        m = df[col].mean()
        if not pd.isna(m):
            label = QUESTION_LABELS.get(col.replace('_num', ''), col)
            means_data.append({'label': label, 'mean': m})
    means_df = pd.DataFrame(means_data).sort_values('mean', ascending=False)
    top_3 = means_df.head(3)['label'].tolist()
    
    return_rate = round(df['q14_return_willingness_label'].value_counts().get
